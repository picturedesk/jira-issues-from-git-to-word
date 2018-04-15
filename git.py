import re, requests, json, sys, subprocess
from docx import Document

# Settings
u = 'm-ammann'
p = open('p.txt')
p = p.readline()
host = 'https://jira.brandleadership.ch'
location = sys.argv[1]

# file handle fh
#input = open('gitcommits_test.txt')
input = subprocess.check_output(["git","--git-dir="+str(location)+".git","log","--pretty=format:'%cd %s'"])
endpoint = '/rest/api/latest/issue/'
url = host + endpoint
document = Document()

for line in input.splitlines():
    # Read line
    line = str(line.strip())
    # Parse for Jira Issue
    number = re.search('[a-zA-Z]+-\d+',line)
    if number:
        target = url + number.group()
        response = requests.get(target, auth=(u,p))

        if response.status_code == requests.codes.ok:
            print(number.group())
            response = response.json()
            #Add to Word pages
            document.add_heading(response["fields"]["summary"], 0)
            document.add_heading('Datum', level=1)
            document.add_paragraph(response["fields"]["created"])
            document.add_heading("Jira Nummer", level=1)
            document.add_paragraph(number.group())
            document.add_heading('Beschreibung', level=1)
            document.add_paragraph(response["fields"]["description"])
            document.add_page_break()
        else:
            print(response)

# Save Word with Jira Issues
document.save('output.docx')